"""
File Upload API Router - Handle PDF/DOCX script uploads
"""
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
import logging
import uuid
import os
from pathlib import Path
import pdfplumber
from docx import Document as DocxDocument

from app.database import get_db
from app.models.database import Project, Document
from app.models.schemas import UploadResponse
from app.config import settings

logger = logging.getLogger(__name__)
router = APIRouter()

# Ensure upload directory exists
UPLOAD_DIR = Path(settings.storage_path) / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


def extract_text_from_pdf(file_path: str) -> tuple[str, int]:
    """Extract text and page count from PDF"""
    try:
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                text += page.extract_text() + "\n"
            return text, len(pdf.pages)
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        raise


def extract_text_from_docx(file_path: str) -> tuple[str, int]:
    """Extract text and page count from DOCX"""
    try:
        doc = DocxDocument(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text, len(doc.paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise


async def validate_project_exists(project_id: str, session: AsyncSession) -> Project:
    """Validate that project exists"""
    result = await session.execute(
        select(Project).where(Project.id == project_id)
    )
    project = result.scalars().first()
    
    if not project:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Project {project_id} not found"
        )
    
    return project


# ============== UPLOAD SCRIPT ==============
@router.post("/{project_id}/upload", response_model=UploadResponse)
async def upload_script(
    project_id: str,
    file: UploadFile = File(...),
    session: AsyncSession = Depends(get_db)
):
    """
    Upload a film script (PDF or DOCX)
    
    **Args:**
    - project_id: UUID of the project
    - file: Script file (PDF or DOCX)
    
    **Returns:** Upload details with document ID
    
    **Supported formats:** .pdf, .docx
    **Max size:** 100 MB
    """
    try:
        # Validate project exists
        await validate_project_exists(project_id, session)
        
        # Validate file format
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in ['.pdf', '.docx']:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Only PDF and DOCX files are supported"
            )
        
        # Validate file size
        file_size_mb = len(await file.read()) / (1024 * 1024)
        await file.seek(0)
        
        if file_size_mb > settings.upload_max_size_mb:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail=f"File size {file_size_mb:.2f}MB exceeds limit of {settings.upload_max_size_mb}MB"
            )
        
        # Save file
        document_id = str(uuid.uuid4())
        file_path = UPLOAD_DIR / f"{document_id}{file_ext}"
        
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        # Extract text
        if file_ext == '.pdf':
            text_content, page_count = extract_text_from_pdf(str(file_path))
            file_format = "pdf"
        else:  # .docx
            text_content, page_count = extract_text_from_docx(str(file_path))
            file_format = "docx"
        
        # Store in database
        document = Document(
            id=document_id,
            project_id=project_id,
            filename=file.filename,
            file_path=str(file_path),
            format=file_format,
            text_content=text_content,
            page_count=page_count
        )
        
        session.add(document)
        await session.commit()
        await session.refresh(document)
        
        logger.info(f"✅ Script uploaded: {document_id} - {file.filename} ({page_count} pages)")
        
        return UploadResponse(
            document_id=document_id,
            filename=file.filename,
            format=file_format,
            page_count=page_count,
            uploaded_at=document.uploaded_at
        )
        
    except HTTPException:
        raise
    except Exception as e:
        await session.rollback()
        logger.error(f"❌ Upload failed: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Upload failed: {str(e)}"
        )


# ============== GET UPLOADED SCRIPT ==============
@router.get("/{project_id}/script/{document_id}")
async def get_script(
    project_id: str,
    document_id: str,
    session: AsyncSession = Depends(get_db)
):
    """
    Get uploaded script details and content
    
    **Args:**
    - project_id: UUID of the project
    - document_id: UUID of the document
    
    **Returns:** Document details with text content
    """
    try:
        result = await session.execute(
            select(Document).where(
                (Document.id == document_id) & 
                (Document.project_id == project_id)
            )
        )
        document = result.scalars().first()
        
        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Document {document_id} not found in project {project_id}"
            )
        
        return {
            "id": document.id,
            "filename": document.filename,
            "format": document.format,
            "page_count": document.page_count,
            "uploaded_at": document.uploaded_at,
            "text_content": document.text_content[:5000],  # First 5000 chars preview
            "text_length": len(document.text_content)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Error fetching script: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error fetching script: {str(e)}"
        )


# ============== LIST PROJECT DOCUMENTS ==============
@router.get("/{project_id}/documents")
async def list_project_documents(
    project_id: str,
    session: AsyncSession = Depends(get_db)
):
    """
    List all uploaded documents for a project
    
    **Args:**
    - project_id: UUID of the project
    
    **Returns:** List of documents
    """
    try:
        # Validate project exists
        await validate_project_exists(project_id, session)
        
        result = await session.execute(
            select(Document).where(Document.project_id == project_id)
        )
        documents = result.scalars().all()
        
        return [
            {
                "id": doc.id,
                "filename": doc.filename,
                "format": doc.format,
                "page_count": doc.page_count,
                "uploaded_at": doc.uploaded_at,
                "text_length": len(doc.text_content) if doc.text_content else 0
            }
            for doc in documents
        ]
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Error listing documents: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error listing documents: {str(e)}"
        )


# ============== DELETE DOCUMENT ==============
@router.delete("/{project_id}/documents/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_document(
    project_id: str,
    document_id: str,
    session: AsyncSession = Depends(get_db)
):
    """
    Delete an uploaded document
    
    **Args:**
    - project_id: UUID of the project
    - document_id: UUID of the document
    
    **Returns:** 204 No Content
    """
    try:
        result = await session.execute(
            select(Document).where(
                (Document.id == document_id) & 
                (Document.project_id == project_id)
            )
        )
        document = result.scalars().first()
        
        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Document {document_id} not found"
            )
        
        # Delete file from storage
        try:
            Path(document.file_path).unlink()
        except:
            logger.warning(f"Could not delete file: {document.file_path}")
        
        # Delete from database
        await session.delete(document)
        await session.commit()
        
        logger.info(f"🗑️ Document deleted: {document_id}")
        
    except HTTPException:
        raise
    except Exception as e:
        await session.rollback()
        logger.error(f"❌ Delete failed: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to delete document: {str(e)}"
        )
